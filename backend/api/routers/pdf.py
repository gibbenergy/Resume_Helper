"""
PDF generation router - resume and cover letter PDFs.
"""

import os
import json
import tempfile
from typing import Dict, Any
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse

from backend.api.dependencies import get_resume_helper
from backend.api.models import GeneratePDFRequest, GenerateDOCXRequest, ResumeData
from backend.core.services import ResumeService as ResumeHelper
from backend.core.workflows.application_workflows import ApplicationWorkflows

router = APIRouter(prefix="/api/pdf", tags=["pdf"])

# Singleton instance to avoid repeated initialization
_app_workflows_instance = None

def get_app_workflows() -> ApplicationWorkflows:
    """Get ApplicationWorkflows instance (singleton to avoid repeated initialization)."""
    global _app_workflows_instance
    if _app_workflows_instance is None:
        _app_workflows_instance = ApplicationWorkflows()
    return _app_workflows_instance


def _build_profile_from_request(resume_dict: dict, resume_helper: ResumeHelper) -> dict:
    """Helper function to build profile dict from request."""
    return resume_helper.build_profile_dict(
        name_prefix=resume_dict.get("personal_info", {}).get("name_prefix", ""),
        email=resume_dict.get("personal_info", {}).get("email", ""),
        full_name=resume_dict.get("personal_info", {}).get("full_name", ""),
        phone=resume_dict.get("personal_info", {}).get("phone", ""),
        current_address=resume_dict.get("personal_info", {}).get("current_address", ""),
        location=resume_dict.get("personal_info", {}).get("location", ""),
        citizenship=resume_dict.get("personal_info", {}).get("citizenship", ""),
        linkedin_url=resume_dict.get("personal_info", {}).get("linkedin_url", ""),
        github_url=resume_dict.get("personal_info", {}).get("github_url", ""),
        portfolio_url=resume_dict.get("personal_info", {}).get("portfolio_url", ""),
        summary=resume_dict.get("personal_info", {}).get("summary", ""),
        education_table=[[
            edu.get("institution", ""),
            edu.get("degree", ""),
            edu.get("field_of_study", ""),
            edu.get("gpa", ""),
            edu.get("start_date", ""),
            edu.get("end_date", ""),
            edu.get("description", "")
        ] for edu in resume_dict.get("education", [])],
        experience_table=[[
            exp.get("company", ""),
            exp.get("position", ""),
            exp.get("location", ""),
            exp.get("start_date", ""),
            exp.get("end_date", ""),
            exp.get("description", ""),
            "\n".join([f"- {a}" for a in exp.get("achievements", [])])
        ] for exp in resume_dict.get("experience", [])],
        skills_table=[[
            skill.get("category", ""),
            skill.get("name", ""),
            skill.get("proficiency", "")
        ] for skill in resume_dict.get("skills", [])],
        projects_table=[[
            proj.get("name", ""),
            proj.get("description", ""),
            proj.get("technologies", ""),
            proj.get("url", ""),
            proj.get("start_date", ""),
            proj.get("end_date", "")
        ] for proj in resume_dict.get("projects", [])],
        certifications_table=[[
            cert.get("name", ""),
            cert.get("issuer", ""),
            cert.get("date_obtained", ""),
            cert.get("credential_id", ""),
            cert.get("url", "")
        ] for cert in resume_dict.get("certifications", [])],
        others_sections_data=resume_dict.get("others", {})
    )


@router.post("/generate-resume")
async def generate_resume_pdf(
    request: GeneratePDFRequest,
    resume_helper: ResumeHelper = Depends(get_resume_helper)
) -> FileResponse:
    """Generate resume PDF."""
    try:
        resume_dict = request.resume_data.dict()
        profile = _build_profile_from_request(resume_dict, resume_helper)
        
        # Generate PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            pdf_path = tmp_file.name
        
        success = resume_helper.resume_gen.generate_pdf(profile, pdf_path)
        
        if not success or not os.path.exists(pdf_path):
            raise HTTPException(status_code=500, detail="Failed to generate PDF")
        
        # Embed metadata if pypdf is available
        try:
            import pypdf
            from pypdf import PdfWriter, PdfReader
            
            with open(pdf_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                pdf_writer = PdfWriter()
                
                for page in pdf_reader.pages:
                    pdf_writer.add_page(page)
                
                metadata = pdf_reader.metadata or {}
                metadata.update({
                    '/ResumeData': json.dumps(profile),
                    '/Creator': 'Resume Helper App',
                    '/ResumeHelperVersion': '1.0'
                })
                pdf_writer.add_metadata(metadata)
                
                with open(pdf_path, 'wb') as output_file:
                    pdf_writer.write(output_file)
        except ImportError:
            pass  # pypdf not available, skip metadata embedding
        except Exception:
            pass  # Metadata embedding failed, continue anyway
        
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename="resume.pdf",
            headers={"Content-Disposition": "attachment; filename=resume.pdf"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating PDF: {str(e)}")


def _generate_docx_resume(profile_data: dict, output_path: str) -> bool:
    """Generate DOCX resume from profile data."""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import base64
        import shutil
        from zipfile import ZipFile, ZIP_DEFLATED
        
        if not profile_data or not isinstance(profile_data, dict):
            return False
        
        doc = Document()
        name_prefix = profile_data.get("name_prefix", "").strip()
        full_name = profile_data.get("full_name", "Resume")
        if name_prefix:
            display_name = f"{name_prefix} {full_name}"
        else:
            display_name = full_name
        
        title = doc.add_heading(display_name, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        contact_para = doc.add_paragraph()
        contact_info = []
        if profile_data.get("email"):
            contact_info.append(str(profile_data["email"]))
        if profile_data.get("phone"):
            contact_info.append(str(profile_data["phone"]))
        if profile_data.get("location") or profile_data.get("current_address"):
            location = profile_data.get("location") or profile_data.get("current_address")
            contact_info.append(str(location))
        
        if contact_info:
            contact_para.add_run(" | ".join(contact_info))
            contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        summary = profile_data.get("summary")
        if summary and str(summary).strip():
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(str(summary))
        
        experience_list = profile_data.get("experience", [])
        if experience_list and isinstance(experience_list, list) and len(experience_list) > 0:
            doc.add_heading("Experience", level=1)
            for exp in experience_list:
                if exp and isinstance(exp, dict):
                    exp_para = doc.add_paragraph()
                    position = exp.get('position', '')
                    company = exp.get('company', '')
                    if position or company:
                        exp_para.add_run(f"{position} at {company}").bold = True
                    
                    start_date = exp.get("start_date", "")
                    end_date = exp.get("end_date", "")
                    if start_date or end_date:
                        exp_para.add_run(f" ({start_date} - {end_date})")
                    
                    description = exp.get("description", "")
                    if description and str(description).strip():
                        desc_para = doc.add_paragraph()
                        desc_para.add_run(str(description).strip())
                    
                    achievements = exp.get("achievements", "")
                    if achievements and isinstance(achievements, list):
                        for achievement in achievements:
                            if achievement and str(achievement).strip():
                                bullet_para = doc.add_paragraph()
                                bullet_para.add_run("• " + str(achievement).strip())
                                bullet_para.paragraph_format.left_indent = Inches(0.5)
                    elif achievements and str(achievements).strip():
                        bullet_para = doc.add_paragraph()
                        bullet_para.add_run("• " + str(achievements).strip())
                        bullet_para.paragraph_format.left_indent = Inches(0.5)
        
        education_list = profile_data.get("education", [])
        if education_list and isinstance(education_list, list) and len(education_list) > 0:
            doc.add_heading("Education", level=1)
            for edu in education_list:
                if edu and isinstance(edu, dict):
                    edu_para = doc.add_paragraph()
                    degree = edu.get('degree', '')
                    institution = edu.get('institution', '')
                    if degree or institution:
                        edu_para.add_run(f"{degree} - {institution}").bold = True
                    
                    start_date = edu.get("start_date", "")
                    end_date = edu.get("end_date", "")
                    if start_date or end_date:
                        edu_para.add_run(f" ({start_date} - {end_date})")
                    
                    description = edu.get("description", "")
                    if description and str(description).strip():
                        desc_para = doc.add_paragraph()
                        desc_para.add_run(str(description).strip())
        
        skills_list = profile_data.get("skills", [])
        if skills_list and isinstance(skills_list, list) and len(skills_list) > 0:
            doc.add_heading("Skills", level=1)
            
            skills_by_category = {}
            uncategorized_skills = []
            
            for skill in skills_list:
                if skill:
                    if isinstance(skill, dict):
                        skill_name = (skill.get("skill", "") or 
                                    skill.get("name", "") or 
                                    skill.get("skill_name", "") or
                                    skill.get("title", ""))
                        category = (skill.get("category", "") or 
                                  skill.get("skill_category", "") or
                                  skill.get("type", ""))
                        
                        if skill_name:
                            if category:
                                if category not in skills_by_category:
                                    skills_by_category[category] = []
                                skills_by_category[category].append(str(skill_name))
                            else:
                                uncategorized_skills.append(str(skill_name))
                    else:
                        uncategorized_skills.append(str(skill))
            
            for category, skills in skills_by_category.items():
                if skills:
                    cat_para = doc.add_paragraph()
                    cat_para.add_run(f"{category}: ").bold = True
                    cat_para.add_run(", ".join(skills))
            
            if uncategorized_skills:
                if skills_by_category:
                    other_para = doc.add_paragraph()
                    other_para.add_run("Other: ").bold = True
                    other_para.add_run(", ".join(uncategorized_skills))
                else:
                    doc.add_paragraph(", ".join(uncategorized_skills))
        
        projects_list = profile_data.get("projects", [])
        if projects_list and isinstance(projects_list, list) and len(projects_list) > 0:
            doc.add_heading("Projects", level=1)
            for project in projects_list:
                if project and isinstance(project, dict):
                    project_para = doc.add_paragraph()
                    title = project.get('title', '') or project.get('name', '')
                    if title:
                        project_para.add_run(str(title)).bold = True
                    
                    description = project.get('description', '')
                    if description and str(description).strip():
                        doc.add_paragraph(str(description))
        
        certifications_list = profile_data.get("certifications", [])
        if certifications_list and isinstance(certifications_list, list) and len(certifications_list) > 0:
            doc.add_heading("Certifications", level=1)
            for cert in certifications_list:
                if cert and isinstance(cert, dict):
                    cert_para = doc.add_paragraph()
                    name = cert.get('name', '') or cert.get('certification', '')
                    issuer = cert.get('issuer', '')
                    if name:
                        if issuer:
                            cert_para.add_run(f"{name} - {issuer}").bold = True
                        else:
                            cert_para.add_run(str(name)).bold = True
        
        others_dict = profile_data.get("others", {})
        if others_dict and isinstance(others_dict, dict):
            for section_name, items in others_dict.items():
                if items and isinstance(items, list) and len(items) > 0:
                    doc.add_heading(section_name, level=1)
                    for item in items:
                        if item and isinstance(item, dict):
                            item_para = doc.add_paragraph()
                            title = item.get('title', '')
                            organization = item.get('organization', '')
                            date = item.get('date', '')
                            location = item.get('location', '')
                            header_parts = []
                            if title:
                                header_parts.append(title)
                            if organization:
                                header_parts.append(organization)
                            if date:
                                header_parts.append(date)
                            if location:
                                header_parts.append(location)
                            
                            if header_parts:
                                item_para.add_run(" | ".join(header_parts)).bold = True
                            description = item.get('description', '')
                            if description:
                                desc_para = doc.add_paragraph(description)
                            url = item.get('url', '')
                            if url and url.strip():
                                url_para = doc.add_paragraph()
                                url_para.add_run("Link: ").bold = True
                                url_para.add_run(url)
        
        doc.save(output_path)
        
        # Embed JSON metadata
        try:
            json_str = json.dumps(profile_data)
            encoded_data = base64.b64encode(json_str.encode('utf-8')).decode('ascii')
            
            custom_props_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
    <property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}" pid="2" name="ResumeHelperData">
        <vt:lpwstr>{encoded_data}</vt:lpwstr>
    </property>
    <property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}" pid="3" name="ResumeHelperVersion">
        <vt:lpwstr>1.0</vt:lpwstr>
    </property>
</Properties>'''
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_path = temp_file.name
            
            shutil.copy2(output_path, temp_path)
            
            with ZipFile(temp_path, 'a', ZIP_DEFLATED, allowZip64=True) as zip_file:
                files_to_keep = [f for f in zip_file.namelist() if f != 'docProps/custom.xml']
                
                new_temp_path = temp_path + '.new'
                with ZipFile(new_temp_path, 'w', ZIP_DEFLATED, allowZip64=True) as new_zip:
                    for item in files_to_keep:
                        data = zip_file.read(item)
                        new_zip.writestr(item, data)
                    
                    new_zip.writestr('docProps/custom.xml', custom_props_xml)
            
            shutil.move(new_temp_path, temp_path)
            shutil.move(temp_path, output_path)
        except Exception:
            pass  # Metadata embedding failed, continue anyway
        
        return True
        
    except ImportError:
        return False
    except Exception:
        return False


@router.post("/generate-resume-docx")
async def generate_resume_docx(
    request: GenerateDOCXRequest,
    resume_helper: ResumeHelper = Depends(get_resume_helper)
) -> FileResponse:
    """Generate resume DOCX."""
    import logging
    logger = logging.getLogger(__name__)
    
    try:
        resume_dict = request.resume_data.dict()
        profile = _build_profile_from_request(resume_dict, resume_helper)
        
        # Check if python-docx is available
        try:
            import docx
        except ImportError as import_err:
            raise HTTPException(
                status_code=500, 
                detail=f"python-docx is not installed. Please install it: pip install python-docx. Error: {str(import_err)}"
            )
        
        # Generate DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            docx_path = tmp_file.name
        
        success = _generate_docx_resume(profile, docx_path)
        
        if not success or not os.path.exists(docx_path):
            raise HTTPException(
                status_code=500, 
                detail="Failed to generate DOCX. Please ensure python-docx is installed: pip install python-docx"
            )
        
        return FileResponse(
            docx_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="resume.docx",
            headers={"Content-Disposition": "attachment; filename=resume.docx"}
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")


@router.post("/generate-cover-letter")
async def generate_cover_letter_pdf(
    request: GeneratePDFRequest,
    resume_helper: ResumeHelper = Depends(get_resume_helper)
) -> FileResponse:
    """Generate cover letter PDF."""
    try:
        from backend.core.infrastructure.generators.cover_letter_generator import generate_cover_letter_pdf
        
        if not request.cover_letter_data:
            raise HTTPException(status_code=400, detail="cover_letter_data is required")
        
        resume_dict = request.resume_data.dict()
        cover_letter_data = request.cover_letter_data
        
        # Flatten personal_info to top level for generator
        personal_info = resume_dict.get("personal_info", {})
        if personal_info:
            resume_dict["full_name"] = personal_info.get("full_name", "")
            resume_dict["name_prefix"] = personal_info.get("name_prefix", "")
            resume_dict["email"] = personal_info.get("email", "")
            resume_dict["phone"] = personal_info.get("phone", "")
            resume_dict["current_address"] = personal_info.get("current_address", "")
            resume_dict["location"] = personal_info.get("location", "")
            resume_dict["linkedin_url"] = personal_info.get("linkedin_url", "")
            resume_dict["github_url"] = personal_info.get("github_url", "")
            resume_dict["portfolio_url"] = personal_info.get("portfolio_url", "")
        
        # Generate PDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            pdf_path = tmp_file.name
        
        # Use the existing cover letter generator
        # Extract cover letter content and metadata
        body_content = cover_letter_data.get('body_content', '') if isinstance(cover_letter_data, dict) else str(cover_letter_data)
        recipient_data = {
            'recipient_name': cover_letter_data.get('recipient_name', '') if isinstance(cover_letter_data, dict) else '',
            'company_name': cover_letter_data.get('company_name', '') if isinstance(cover_letter_data, dict) else '',
            'company_address': cover_letter_data.get('company_address', '') if isinstance(cover_letter_data, dict) else '',
        }
        job_position = cover_letter_data.get('job_position', '') if isinstance(cover_letter_data, dict) else ''
        company_name = cover_letter_data.get('company_name', '') if isinstance(cover_letter_data, dict) else ''
        letter_title = cover_letter_data.get('letter_title', '') if isinstance(cover_letter_data, dict) else ''
        recipient_greeting = cover_letter_data.get('recipient_greeting', '') if isinstance(cover_letter_data, dict) else ''
        
        # Get job analysis data if available
        job_analysis = request.job_analysis_data or {}
        if job_analysis:
            job_position = job_analysis.get('position_title', job_position)
            company_name = job_analysis.get('company_name', company_name)
        
        pdf_path = generate_cover_letter_pdf(
            candidate_data=resume_dict,
            cover_letter_content=body_content,
            recipient_data=recipient_data,
            output_path=pdf_path,
            temp_dir=resume_helper.resume_gen.temp_dir,
            job_position=job_position,
            company_name=company_name,
            letter_title=letter_title,
            recipient_greeting=recipient_greeting,
        )
        
        success = pdf_path and os.path.exists(pdf_path)
        
        if not success:
            raise HTTPException(status_code=500, detail="Failed to generate cover letter PDF")
        
        import datetime as dt
        import re
        
        def _slug(text: str, default: str = "document") -> str:
            """Convert arbitrary text to filesystem-safe slug."""
            text = re.sub(r"[^A-Za-z0-9]+", "_", text or "")
            return text.strip("_").lower() or default
        timestamp = dt.datetime.now().strftime("%Y%m%d")
        filename = f"{_slug(company_name or 'company')}_cover_letter_{timestamp}.pdf"
        
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating cover letter PDF: {str(e)}")


@router.post("/generate-job-analysis")
async def generate_job_analysis_pdf(
    request: Dict[str, Any],
    resume_helper: ResumeHelper = Depends(get_resume_helper)
) -> FileResponse:
    """Generate job analysis PDF."""
    try:
        from backend.core.infrastructure.generators.analysis_pdf_generator import generate_job_analysis_pdf
        import datetime as dt
        import re
        
        def _slug(text: str, default: str = "document") -> str:
            """Convert arbitrary text to filesystem-safe slug."""
            text = re.sub(r"[^A-Za-z0-9]+", "_", text or "")
            return text.strip("_").lower() or default
        
        if not request.get("analysis_data"):
            raise HTTPException(status_code=400, detail="analysis_data is required")
        
        analysis_data = request["analysis_data"]
        
        # Format analysis as markdown
        formatted_content = ""
        if isinstance(analysis_data, dict):
            # Format the analysis data as markdown with proper line breaks
            # DON'T include position_title in markdown - it's already in the template header
            if "company_name" in analysis_data:
                formatted_content += f"**Company:**\n\n{analysis_data.get('company_name', 'Not Specified')}\n\n"
            if "location" in analysis_data:
                formatted_content += f"**Location:**\n\n{analysis_data.get('location', 'Not Specified')}\n\n"
            if "match_score" in analysis_data:
                formatted_content += f"**Match Score:**\n\n{analysis_data.get('match_score', 0)}%\n\n"
            if "match_summary" in analysis_data:
                formatted_content += f"**Summary:**\n\n{analysis_data.get('match_summary', '')}\n\n"
            
            # Add all other fields - position_title is already in template header
            for key, value in analysis_data.items():
                if key not in ["position_title", "company_name", "location", "match_score", "match_summary"]:
                    if isinstance(value, list) and value:
                        formatted_content += f"**{key.replace('_', ' ').title()}:**\n\n"
                        for item in value:
                            formatted_content += f"- {item}\n"
                        formatted_content += "\n"
                    elif isinstance(value, dict) and value:
                        formatted_content += f"**{key.replace('_', ' ').title()}:**\n\n"
                        for k, v in value.items():
                            formatted_content += f"- **{k.replace('_', ' ').title()}:** {v}\n"
                        formatted_content += "\n"
                    elif value:  # Only add if value is not empty
                        formatted_content += f"**{key.replace('_', ' ').title()}:**\n\n{value}\n\n"
        else:
            formatted_content = str(analysis_data)
        
        # Generate PDF
        temp_dir = resume_helper.resume_gen.temp_dir
        os.makedirs(temp_dir, exist_ok=True)
        
        company = analysis_data.get("company_name", "") if isinstance(analysis_data, dict) else "company"
        timestamp = dt.datetime.now().strftime("%Y%m%d")
        filename = f"{_slug(company)}_job_analysis_{timestamp}.pdf"
        pdf_path = os.path.join(temp_dir, filename)
        
        pdf_path = generate_job_analysis_pdf(
            analysis_content=formatted_content,
            company_name=company,
            job_position=analysis_data.get("position_title", "") if isinstance(analysis_data, dict) else "",
            output_path=pdf_path,
            temp_dir=temp_dir
        )
        
        if not pdf_path or not os.path.exists(pdf_path):
            raise HTTPException(status_code=500, detail="Failed to generate job analysis PDF")
        
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating job analysis PDF: {str(e)}")


@router.post("/generate-tailored-resume")
async def generate_tailored_resume_pdf(
    request: Dict[str, Any],
    resume_helper: ResumeHelper = Depends(get_resume_helper)
) -> FileResponse:
    """Generate tailored resume PDF with metadata embedding."""
    try:
        import datetime as dt
        import re
        
        def _slug(text: str, default: str = "document") -> str:
            """Convert arbitrary text to filesystem-safe slug."""
            text = re.sub(r"[^A-Za-z0-9]+", "_", text or "")
            return text.strip("_").lower() or default
        
        if not request.get("tailored_resume_data"):
            raise HTTPException(status_code=400, detail="tailored_resume_data is required")
        
        tailored_resume = request["tailored_resume_data"]
        
        # Flatten personal_info to top level for generator
        if isinstance(tailored_resume, dict):
            personal_info = tailored_resume.get("personal_info", {})
            if personal_info:
                # Only set if not already present (tailored resume may already have flattened fields)
                if "full_name" not in tailored_resume:
                    tailored_resume["full_name"] = personal_info.get("full_name", "")
                if "name_prefix" not in tailored_resume:
                    tailored_resume["name_prefix"] = personal_info.get("name_prefix", "")
                if "email" not in tailored_resume:
                    tailored_resume["email"] = personal_info.get("email", "")
                if "phone" not in tailored_resume:
                    tailored_resume["phone"] = personal_info.get("phone", "")
                if "current_address" not in tailored_resume:
                    tailored_resume["current_address"] = personal_info.get("current_address", "")
                if "location" not in tailored_resume:
                    tailored_resume["location"] = personal_info.get("location", "")
                if "linkedin_url" not in tailored_resume:
                    tailored_resume["linkedin_url"] = personal_info.get("linkedin_url", "")
                if "github_url" not in tailored_resume:
                    tailored_resume["github_url"] = personal_info.get("github_url", "")
                if "portfolio_url" not in tailored_resume:
                    tailored_resume["portfolio_url"] = personal_info.get("portfolio_url", "")
        
        # Generate PDF
        temp_dir = resume_helper.resume_gen.temp_dir
        os.makedirs(temp_dir, exist_ok=True)
        
        company = tailored_resume.get("company_name", "") if isinstance(tailored_resume, dict) else "company"
        timestamp = dt.datetime.now().strftime("%Y%m%d")
        filename = f"{_slug(company)}_resume_{timestamp}.pdf"
        pdf_path = os.path.join(temp_dir, filename)
        
        # Generate PDF using resume generator
        success = resume_helper.resume_gen.generate_pdf(tailored_resume, pdf_path)
        
        if not success or not os.path.exists(pdf_path):
            raise HTTPException(status_code=500, detail="Failed to generate tailored resume PDF")
        
        # Embed metadata if pypdf is available
        try:
            import pypdf
            from pypdf import PdfWriter, PdfReader
            
            with open(pdf_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                pdf_writer = PdfWriter()
                
                for page in pdf_reader.pages:
                    pdf_writer.add_page(page)
                
                metadata = pdf_reader.metadata or {}
                metadata.update({
                    '/ResumeData': json.dumps(tailored_resume),
                    '/Creator': 'Resume Helper App - AI Resume Helper',
                    '/ResumeHelperVersion': '1.0',
                    '/AIGenerated': 'true',
                    '/Source': 'ai_resume_helper'
                })
                pdf_writer.add_metadata(metadata)
                
                with open(pdf_path, 'wb') as output_file:
                    pdf_writer.write(output_file)
        except ImportError:
            pass  # pypdf not available, skip metadata embedding
        except Exception:
            pass  # Metadata embedding failed, continue anyway
        
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating tailored resume PDF: {str(e)}")


@router.post("/generate-suggestions")
async def generate_suggestions_pdf(
    request: Dict[str, Any],
    resume_helper: ResumeHelper = Depends(get_resume_helper)
) -> FileResponse:
    """Generate improvement suggestions PDF."""
    try:
        from backend.core.infrastructure.generators.analysis_pdf_generator import generate_improvement_suggestions_pdf
        import datetime as dt
        import re
        
        def _slug(text: str, default: str = "document") -> str:
            """Convert arbitrary text to filesystem-safe slug."""
            text = re.sub(r"[^A-Za-z0-9]+", "_", text or "")
            return text.strip("_").lower() or default
        
        if not request.get("suggestions_content"):
            raise HTTPException(status_code=400, detail="suggestions_content is required")
        
        suggestions_content = request["suggestions_content"]
        full_name = request.get("full_name", "")
        company_name = request.get("company_name", "company")
        job_position = request.get("job_position", "")
        
        # Generate PDF
        temp_dir = resume_helper.resume_gen.temp_dir
        os.makedirs(temp_dir, exist_ok=True)
        
        timestamp = dt.datetime.now().strftime("%Y%m%d")
        filename = f"{_slug(company_name)}_suggestions_{timestamp}.pdf"
        pdf_path = os.path.join(temp_dir, filename)
        
        pdf_path = generate_improvement_suggestions_pdf(
            suggestions_content=suggestions_content,
            full_name=full_name,
            company_name=company_name,
            job_position=job_position,
            output_path=pdf_path,
            temp_dir=temp_dir
        )
        
        if not pdf_path or not os.path.exists(pdf_path):
            raise HTTPException(status_code=500, detail="Failed to generate suggestions PDF")
        
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=filename,
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating suggestions PDF: {str(e)}")


@router.get("/job-description/{app_id}")
async def get_job_description_pdf(
    app_id: str,
    workflows: ApplicationWorkflows = Depends(get_app_workflows),
    resume_helper: ResumeHelper = Depends(get_resume_helper)
) -> FileResponse:
    """Get job description PDF for an application (generate if not exists)."""
    try:
        from backend.core.infrastructure.generators.analysis_pdf_generator import generate_job_analysis_pdf
        import datetime as dt
        import re
        import hashlib
        
        def _slug(text: str, default: str = "document") -> str:
            """Convert arbitrary text to filesystem-safe slug."""
            text = re.sub(r"[^A-Za-z0-9]+", "_", text or "")
            return text.strip("_").lower() or default
        
        # Get application data
        app = workflows.get_application(app_id)
        
        if not app:
            raise HTTPException(status_code=404, detail=f"Application {app_id} not found")
        
        if not app.get("description"):
            raise HTTPException(status_code=404, detail="No job description available")
        
        description = app["description"]
        company_name = app.get("company", "company")
        job_position = app.get("position", "")
        
        # Store PDFs in a dedicated directory
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        pdfs_dir = os.path.join(project_root, "data", "job_description_pdfs")
        os.makedirs(pdfs_dir, exist_ok=True)
        
        # Use app_id as filename - one PDF per application (overwrites on description change)
        filename = f"{app_id}_job_description.pdf"
        pdf_path = os.path.join(pdfs_dir, filename)
        
        # Create hash of current description
        desc_hash = hashlib.md5(description.encode()).hexdigest()[:8]
        
        # Check if PDF exists and if description has changed
        # We'll regenerate if PDF doesn't exist or if description hash doesn't match
        needs_regeneration = True
        hash_file = os.path.join(pdfs_dir, f"{app_id}_hash.txt")
        
        if os.path.exists(pdf_path) and os.path.exists(hash_file):
            try:
                # Check if there's a hash file to compare
                with open(hash_file, 'r') as f:
                    stored_hash = f.read().strip()
                if stored_hash == desc_hash:
                    needs_regeneration = False
            except (IOError, OSError) as e:
                # If hash file read fails, regenerate PDF
                needs_regeneration = True
        
        # Generate PDF only if needed
        if needs_regeneration:
            # Format description as markdown for PDF
            formatted_content = f"# Job Description\n\n"
            if company_name:
                formatted_content += f"**Company:** {company_name}\n\n"
            if job_position:
                formatted_content += f"**Position:** {job_position}\n\n"
            formatted_content += f"{description}\n"
            
            # Generate PDF
            temp_dir = resume_helper.resume_gen.temp_dir
            pdf_path = generate_job_analysis_pdf(
                analysis_content=formatted_content,
                company_name=company_name,
                job_position=job_position,
                output_path=pdf_path,
                temp_dir=temp_dir
            )
            
            if not pdf_path or not os.path.exists(pdf_path):
                raise HTTPException(status_code=500, detail="Failed to generate job description PDF")
            
            # Store hash for future comparison
            try:
                with open(hash_file, 'w') as f:
                    f.write(desc_hash)
            except (IOError, OSError) as e:
                # Log error but don't fail - hash file is optional
                pass
        
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=filename,
            headers={"Content-Disposition": f"inline; filename={filename}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting job description PDF: {str(e)}")


@router.post("/generate-job-description")
async def generate_job_description_pdf(
    request: Dict[str, Any],
    resume_helper: ResumeHelper = Depends(get_resume_helper)
) -> FileResponse:
    """Generate job description PDF."""
    try:
        from backend.core.infrastructure.generators.analysis_pdf_generator import generate_job_analysis_pdf
        import datetime as dt
        import re
        
        def _slug(text: str, default: str = "document") -> str:
            """Convert arbitrary text to filesystem-safe slug."""
            text = re.sub(r"[^A-Za-z0-9]+", "_", text or "")
            return text.strip("_").lower() or default
        
        if not request.get("description"):
            raise HTTPException(status_code=400, detail="description is required")
        
        description = request["description"]
        company_name = request.get("company_name", "company")
        job_position = request.get("job_position", "")
        
        # Format description as markdown for PDF
        formatted_content = f"# Job Description\n\n"
        if company_name:
            formatted_content += f"**Company:** {company_name}\n\n"
        if job_position:
            formatted_content += f"**Position:** {job_position}\n\n"
        formatted_content += f"{description}\n"
        
        # Generate PDF
        temp_dir = resume_helper.resume_gen.temp_dir
        os.makedirs(temp_dir, exist_ok=True)
        
        timestamp = dt.datetime.now().strftime("%Y%m%d")
        filename = f"{_slug(company_name)}_job_description_{timestamp}.pdf"
        pdf_path = os.path.join(temp_dir, filename)
        
        pdf_path = generate_job_analysis_pdf(
            analysis_content=formatted_content,
            company_name=company_name,
            job_position=job_position,
            output_path=pdf_path,
            temp_dir=temp_dir
        )
        
        if not pdf_path or not os.path.exists(pdf_path):
            raise HTTPException(status_code=500, detail="Failed to generate job description PDF")
        
        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=filename,
            headers={"Content-Disposition": f"inline; filename={filename}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating job description PDF: {str(e)}") 
