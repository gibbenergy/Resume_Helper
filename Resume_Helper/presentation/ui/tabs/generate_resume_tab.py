from types import SimpleNamespace
import os, json, uuid
import gradio as gr
try:
    import pypdf
except ImportError:
    pypdf = None


def create_generate_resume_tab(resume_helper, all_tabs_components=None):
    if all_tabs_components is None:
        all_tabs_components = {}

    with gr.Tab("Import & Export") as tab:
        with gr.Group():
            gr.Markdown("## 📥 **Import Resume**")
            gr.Markdown("*Import your existing resume data from supported formats*")
            
            with gr.Row():
                with gr.Column(scale=3):
                    file_import = gr.File(
                        label="📂 Select Resume File", 
                        file_types=[".json", ".pdf", ".docx", ".doc"], 
                        type="filepath",
                        height=120
                    )
                with gr.Column(scale=2):
                    gr.Markdown("**Supported Formats:**")
                    gr.Markdown("• **JSON** - Resume Helper exports")
                    gr.Markdown("• **PDF** - Resume Helper generated")  
                    gr.Markdown("• **DOCX/DOC** - Microsoft Word files")
            
            import_status = gr.Textbox(
                label="Import Status", 
                interactive=False, 
                lines=1,
                value="Ready to import resume file..."
            )

        with gr.Group():
            gr.Markdown("## 📤 **Export Resume**")
            gr.Markdown("*Generate your resume in your preferred format*")
            
            export_format = gr.Radio(
                label="📋 Select Export Format",
                choices=[
                    ("📄 PDF Document", "pdf"),
                    ("💾 JSON Data", "json"), 
                    ("📝 Word Document (.docx)", "docx")
                ],
                value="pdf",
                interactive=True
            )
            
            with gr.Row():
                generate_btn = gr.Button(
                    "🚀 Generate Resume", 
                    variant="primary", 
                    size="lg",
                    scale=1
                )
                download_btn = gr.DownloadButton(
                    "⬇️ Download Resume", 
                    variant="secondary", 
                    interactive=False,
                    size="lg",
                    scale=1
                )
                
            export_status = gr.Textbox(
                label="Export Status", 
                interactive=False, 
                lines=1,
                value="Select format and click Generate to create your resume..."
            )

        def get_empty_form_data():
            """Return properly structured empty form data matching expected outputs.
            
            This function dynamically builds the empty data structure using UIConstants
            to ensure it adapts automatically to schema changes.
            """
            from utils.constants import UIConstants
            
            result = [
                "Ready to import profile",
                *[""] * UIConstants.PERSONAL_INFO_FIELDS,
                *[""] * UIConstants.EDUCATION_INPUT_FIELDS,
                [],
                *[""] * UIConstants.WORK_INPUT_FIELDS,
                [],
                *[""] * UIConstants.SKILLS_INPUT_FIELDS,
                [],
                *[""] * UIConstants.PROJECTS_INPUT_FIELDS,
                [],
                *[""] * UIConstants.CERTIFICATIONS_INPUT_FIELDS,
                [],
                {},
                [],
                gr.update(choices=[], value=None)
            ]
            
            expected_count = UIConstants.TOTAL_OUTPUT_COUNT + 1
            actual_count = len(result)
            
            if actual_count != expected_count:
                from utils.logging_helpers import log_error
                log_error(f"get_empty_form_data count mismatch: expected {expected_count}, got {actual_count}")
                
                if actual_count < expected_count:
                    result.extend([""] * (expected_count - actual_count))
                elif actual_count > expected_count:
                    result = result[:expected_count]
            
            return result

        def extract_json_from_docx(docx_path):
            """Extract embedded JSON metadata from DOCX file."""
            try:
                import json
                import base64
                import tempfile
                from zipfile import ZipFile
                import xml.etree.ElementTree as ET
                
                with ZipFile(docx_path, 'r') as zip_file:
                    if 'docProps/custom.xml' in zip_file.namelist():
                        custom_props = zip_file.read('docProps/custom.xml')
                        root = ET.fromstring(custom_props)
                        
                        for prop in root.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/custom-properties}property'):
                            prop_name = prop.get('name')
                            
                            if prop_name == 'ResumeHelperData':
                                text_elem = prop.find('.//{http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes}lpwstr')
                                if text_elem is not None and text_elem.text:
                                    try:
                                        encoded_data = text_elem.text
                                        json_str = base64.b64decode(encoded_data).decode('utf-8')
                                        profile_data = json.loads(json_str)
                                        
                                        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as temp_file:
                                            json.dump(profile_data, temp_file, indent=2)
                                            temp_path = temp_file.name
                                        
                                        try:
                                            result = resume_helper.load_from_json(temp_path)
                                            return result, profile_data
                                        finally:
                                            os.unlink(temp_path)
                                    except Exception:
                                        pass
                
                try:
                    from docx import Document
                    doc = Document(docx_path)
                    text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                except Exception:
                    pass
                
                return None, None
                
            except ImportError as e:
                from utils.logging_helpers import log_error
                log_error(f"Import error in extract_json_from_docx: {e}")
                return None, None
            except Exception as e:
                from utils.logging_helpers import log_error
                log_error(f"General error in extract_json_from_docx: {e}")
                return None, None

        def embed_json_metadata_docx(docx_path, profile_data):
            """Embed JSON data as custom properties in DOCX file."""
            try:
                import json
                import base64
                import tempfile
                import shutil
                from zipfile import ZipFile, ZIP_DEFLATED
                
                json_str = json.dumps(profile_data)
                encoded_data = base64.b64encode(json_str.encode('utf-8')).decode('ascii')
                
                custom_props_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">
    <property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}" pid="2" name="ResumeHelperData">
        <vt:lpwstr>{encoded_data}</vt:lpwstr>
    </property>
    <property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}" pid="3" name="ResumeHelperVersion">
        <vt:lpwstr>1.0</vt:lpwstr>
    </property>
</Properties>'''
                
                with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                    temp_path = temp_file.name
                
                shutil.copy2(docx_path, temp_path)
                
                with ZipFile(temp_path, 'a', ZIP_DEFLATED, allowZip64=True) as zip_file:
                    files_to_keep = [f for f in zip_file.namelist() if f != 'docProps/custom.xml']
                    
                    new_temp_path = temp_path + '.new'
                    with ZipFile(new_temp_path, 'w', ZIP_DEFLATED, allowZip64=True) as new_zip:
                        for item in files_to_keep:
                            data = zip_file.read(item)
                            new_zip.writestr(item, data)
                        
                        new_zip.writestr('docProps/custom.xml', custom_props_xml)
                
                shutil.move(new_temp_path, temp_path)
                shutil.move(temp_path, docx_path)
                
            except Exception as e:
                from utils.logging_helpers import log_error
                log_error(f"Error embedding metadata in DOCX: {e}")

        def auto_load_file(uploaded_file):
            """Enhanced auto-load with DOCX support."""
            if (uploaded_file is None or 
                uploaded_file == "" or
                uploaded_file == [] or
                not uploaded_file or 
                not hasattr(uploaded_file, 'name') or 
                uploaded_file.name is None or 
                uploaded_file.name == "" or
                not uploaded_file.name or 
                not uploaded_file.name.strip() or
                uploaded_file.name.strip() == ""):
                return get_empty_form_data(), None
                
            try:
                file_path = uploaded_file.name.strip()
                if not file_path or file_path == "":
                    return get_empty_form_data(), None
                if not os.path.exists(file_path):
                    empty_data = get_empty_form_data()
                    empty_data[0] = f"❌ File not found: {file_path}"
                    return empty_data, None
                
                filename = os.path.basename(file_path)
                file_ext = os.path.splitext(filename)[1].lower()
                original_json_data = None
                
                if file_ext == '.json':
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            original_json_data = json.load(f)
                        
                        profile_vals = resume_helper.load_from_json(file_path)
                        status_msg = f"✅ Successfully imported JSON resume: {filename}"
                    except (ValueError, FileNotFoundError) as e:
                        empty_data = get_empty_form_data()
                        empty_data[0] = f"❌ JSON import error: {str(e)}"
                        return empty_data, None
                        
                elif file_ext == '.pdf':
                    try:
                        import pypdf
                        with open(file_path, 'rb') as file:
                            pdf_reader = pypdf.PdfReader(file)
                            metadata = pdf_reader.metadata
                            
                            if metadata and '/ResumeData' in metadata:
                                json_data_str = str(metadata['/ResumeData'])
                                original_json_data = json.loads(json_data_str)
                    except Exception:
                        pass
                    
                    profile_vals = extract_json_from_pdf(file_path)
                    if profile_vals is None:
                        empty_data = get_empty_form_data()
                        empty_data[0] = f"❌ No resume data found in PDF: {filename}"
                        return empty_data, None
                    status_msg = f"✅ Successfully imported PDF resume: {filename}"
                    
                elif file_ext in ['.docx', '.doc']:
                    try:
                        profile_vals, original_json_data = extract_json_from_docx(file_path)
                        if profile_vals is None:
                            empty_data = get_empty_form_data()
                            empty_data[0] = f"❌ Could not extract resume data from Word document: {filename}"
                            return empty_data, None
                        status_msg = f"✅ Successfully imported Word document: {filename}"
                    except Exception:
                        empty_data = get_empty_form_data()
                        empty_data[0] = f"❌ Error processing Word document: {filename}"
                        return empty_data, None
                        
                else:
                    empty_data = get_empty_form_data()
                    empty_data[0] = f"❌ Unsupported file type: {file_ext} (Supported: .json, .pdf, .docx, .doc)"
                    return empty_data, None
                
                from utils.constants import UIConstants
                from utils.logging_helpers import StandardLogger
                import uuid
                
                request_id = str(uuid.uuid4())
                StandardLogger.log_operation_start("process_file_import", request_id, 
                                                 file_type=file_ext, filename=filename)
                
                expected_profile_vals = UIConstants.FORM_BASE_COUNT + 1
                if len(profile_vals) != expected_profile_vals:
                    StandardLogger.log_operation_warning("process_file_import", request_id,
                                                        f"Unexpected profile_vals count: {len(profile_vals)}, expected {expected_profile_vals}")
                
                others_raw_data = profile_vals[-1] if len(profile_vals) >= expected_profile_vals else {}
                
                adjusted_vals_without_others = profile_vals[:-1]
                
                if others_raw_data and isinstance(others_raw_data, dict):
                    others_table = []
                    for section_name, items in others_raw_data.items():
                        if isinstance(items, list):
                            for item in items:
                                if isinstance(item, dict):
                                    others_table.append([
                                        section_name,
                                        item.get("title", ""),
                                        item.get("organization", ""),
                                        item.get("date", ""),
                                        item.get("location", ""),
                                        item.get("description", ""),
                                        item.get("url", "")
                                    ])
                    
                    others_choices = list(others_raw_data.keys())
                    selected_section = others_choices[0] if others_choices else None
                    result = [status_msg] + adjusted_vals_without_others + [
                        others_raw_data,
                        others_table,
                        gr.update(choices=others_choices, value=selected_section)
                    ]
                else:
                    result = [status_msg] + adjusted_vals_without_others + [
                        {},
                        [],
                        gr.update(choices=[], value=None)
                    ]
                
                expected_count = UIConstants.TOTAL_OUTPUT_COUNT + 1
                actual_count = len(result)
                
                if actual_count != expected_count:
                    error_msg = f"Output count mismatch: expected {expected_count}, got {actual_count}"
                    StandardLogger.log_operation_error("process_file_import", request_id, 
                                                     Exception(error_msg))
                    empty_data = get_empty_form_data()
                    empty_data[0] = f"❌ {error_msg}"
                    return empty_data, None
                
                StandardLogger.log_operation_success("process_file_import", request_id, 
                                                   output_count=actual_count)
                
                return result, original_json_data
                
            except Exception:
                empty_data = get_empty_form_data()
                empty_data[0] = "❌ Error loading file"
                return empty_data, None

        def extract_json_from_pdf(pdf_path):
            """Extract embedded JSON metadata from PDF file."""
            try:
                import pypdf
                with open(pdf_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    metadata = pdf_reader.metadata
                    
                    if metadata and '/ResumeData' in metadata:
                        json_data = str(metadata['/ResumeData'])
                        profile_data = json.loads(json_data)
                        
                        import tempfile
                        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False) as temp_file:
                            json.dump(profile_data, temp_file, indent=2)
                            temp_path = temp_file.name
                        
                        try:
                            result = resume_helper.load_from_json(temp_path)
                            return result
                        finally:
                            os.unlink(temp_path)
                    else:
                        return None
            except ImportError:
                return None
            except Exception:
                return None

        def embed_json_metadata_pdf(pdf_path, profile_data):
            """Embed JSON data as metadata in PDF file."""
            try:
                import pypdf
                from pypdf import PdfWriter, PdfReader
                
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PdfReader(file)
                    pdf_writer = PdfWriter()
                    
                    for page in pdf_reader.pages:
                        pdf_writer.add_page(page)
                    
                    metadata = pdf_reader.metadata or {}
                    metadata.update({
                        '/ResumeData': json.dumps(profile_data),
                        '/Creator': 'Resume Helper App',
                        '/ResumeHelperVersion': '1.0'
                    })
                    pdf_writer.add_metadata(metadata)
                    
                    with open(pdf_path, 'wb') as output_file:
                        pdf_writer.write(output_file)
                        
            except ImportError:
                pass
            except Exception:
                pass

        def generate_docx_resume(profile_data, output_path):
            try:
                from docx import Document
                from docx.shared import Inches
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
                
                if not profile_data or not isinstance(profile_data, dict):
                    return False
                
                doc = Document()
                name_prefix = profile_data.get("name_prefix", "").strip()
                full_name = profile_data.get("full_name", "Resume")
                if name_prefix:
                    display_name = f"{name_prefix} {full_name}"
                else:
                    display_name = full_name
                
                title = doc.add_heading(display_name, 0)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                contact_para = doc.add_paragraph()
                contact_info = []
                if profile_data.get("email"):
                    contact_info.append(str(profile_data["email"]))
                if profile_data.get("phone"):
                    contact_info.append(str(profile_data["phone"]))
                if profile_data.get("location") or profile_data.get("current_address"):
                    location = profile_data.get("location") or profile_data.get("current_address")
                    contact_info.append(str(location))
                
                if contact_info:
                    contact_para.add_run(" | ".join(contact_info))
                    contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                summary = profile_data.get("summary")
                if summary and str(summary).strip():
                    doc.add_heading("Summary", level=1)
                    doc.add_paragraph(str(summary))
                experience_list = profile_data.get("experience", [])
                if experience_list and isinstance(experience_list, list) and len(experience_list) > 0:
                    doc.add_heading("Experience", level=1)
                    for exp in experience_list:
                        if exp and isinstance(exp, dict):
                            exp_para = doc.add_paragraph()
                            position = exp.get('position', '')
                            company = exp.get('company', '')
                            if position or company:
                                exp_para.add_run(f"{position} at {company}").bold = True
                            
                            start_date = exp.get("start_date", "")
                            end_date = exp.get("end_date", "")
                            if start_date or end_date:
                                exp_para.add_run(f" ({start_date} - {end_date})")
                            
                            description = exp.get("description", "")
                            if description and str(description).strip():
                                desc_para = doc.add_paragraph()
                                desc_para.add_run(str(description).strip())
                            
                            achievements = exp.get("achievements", "")
                            if achievements and isinstance(achievements, list):
                                for achievement in achievements:
                                    if achievement and str(achievement).strip():
                                        bullet_para = doc.add_paragraph()
                                        bullet_para.add_run("• " + str(achievement).strip())
                                        bullet_para.paragraph_format.left_indent = Inches(0.5)
                            elif achievements and str(achievements).strip():
                                bullet_para = doc.add_paragraph()
                                bullet_para.add_run("• " + str(achievements).strip())
                                bullet_para.paragraph_format.left_indent = Inches(0.5)  
                education_list = profile_data.get("education", [])
                if education_list and isinstance(education_list, list) and len(education_list) > 0:
                    doc.add_heading("Education", level=1)
                    for edu in education_list:
                        if edu and isinstance(edu, dict):
                            edu_para = doc.add_paragraph()
                            degree = edu.get('degree', '')
                            institution = edu.get('institution', '')
                            if degree or institution:
                                edu_para.add_run(f"{degree} - {institution}").bold = True
                            
                            start_date = edu.get("start_date", "")
                            end_date = edu.get("end_date", "")
                            if start_date or end_date:
                                edu_para.add_run(f" ({start_date} - {end_date})")
                            
                            description = edu.get("description", "")
                            if description and str(description).strip():
                                desc_para = doc.add_paragraph()
                                desc_para.add_run(str(description).strip())
                
                skills_list = profile_data.get("skills", [])
                if skills_list and isinstance(skills_list, list) and len(skills_list) > 0:
                    doc.add_heading("Skills", level=1)
                    
                    skills_by_category = {}
                    uncategorized_skills = []
                    
                    for skill in skills_list:
                        if skill:
                            if isinstance(skill, dict):
                                skill_name = (skill.get("skill", "") or 
                                            skill.get("name", "") or 
                                            skill.get("skill_name", "") or
                                            skill.get("title", ""))
                                category = (skill.get("category", "") or 
                                          skill.get("skill_category", "") or
                                          skill.get("type", ""))
                                
                                if skill_name:
                                    if category:
                                        if category not in skills_by_category:
                                            skills_by_category[category] = []
                                        skills_by_category[category].append(str(skill_name))
                                    else:
                                        uncategorized_skills.append(str(skill_name))
                            else:
                                uncategorized_skills.append(str(skill))
                    
                    for category, skills in skills_by_category.items():
                        if skills:
                            cat_para = doc.add_paragraph()
                            cat_para.add_run(f"{category}: ").bold = True
                            cat_para.add_run(", ".join(skills))
                    
                    if uncategorized_skills:
                        if skills_by_category:
                            other_para = doc.add_paragraph()
                            other_para.add_run("Other: ").bold = True
                            other_para.add_run(", ".join(uncategorized_skills))
                        else:
                            doc.add_paragraph(", ".join(uncategorized_skills))
                projects_list = profile_data.get("projects", [])
                if projects_list and isinstance(projects_list, list) and len(projects_list) > 0:
                    doc.add_heading("Projects", level=1)
                    for project in projects_list:
                        if project and isinstance(project, dict):
                            project_para = doc.add_paragraph()
                            title = project.get('title', '') or project.get('name', '')
                            if title:
                                project_para.add_run(str(title)).bold = True
                            
                            description = project.get('description', '')
                            if description and str(description).strip():
                                doc.add_paragraph(str(description))
                
                certifications_list = profile_data.get("certifications", [])
                if certifications_list and isinstance(certifications_list, list) and len(certifications_list) > 0:
                    doc.add_heading("Certifications", level=1)
                    for cert in certifications_list:
                        if cert and isinstance(cert, dict):
                            cert_para = doc.add_paragraph()
                            name = cert.get('name', '') or cert.get('certification', '')
                            issuer = cert.get('issuer', '')
                            if name:
                                if issuer:
                                    cert_para.add_run(f"{name} - {issuer}").bold = True
                                else:
                                    cert_para.add_run(str(name)).bold = True
                others_dict = profile_data.get("others", {})
                if others_dict and isinstance(others_dict, dict):
                    for section_name, items in others_dict.items():
                        if items and isinstance(items, list) and len(items) > 0:
                            doc.add_heading(section_name, level=1)
                            for item in items:
                                if item and isinstance(item, dict):
                                    item_para = doc.add_paragraph()
                                    title = item.get('title', '')
                                    organization = item.get('organization', '')
                                    date = item.get('date', '')
                                    location = item.get('location', '')
                                    header_parts = []
                                    if title:
                                        header_parts.append(title)
                                    if organization:
                                        header_parts.append(organization)
                                    if date:
                                        header_parts.append(date)
                                    if location:
                                        header_parts.append(location)
                                    
                                    if header_parts:
                                        item_para.add_run(" | ".join(header_parts)).bold = True
                                    description = item.get('description', '')
                                    if description:
                                        desc_para = doc.add_paragraph(description)
                                    url = item.get('url', '')
                                    if url and url.strip():
                                        url_para = doc.add_paragraph()
                                        url_para.add_run("Link: ").bold = True
                                        url_para.add_run(url)
                
                doc.save(output_path)
                embed_json_metadata_docx(output_path, profile_data)
                return True
                
            except ImportError:
                return False
            except Exception:
                return False

        def generate_resume(format_choice):
            try:
                temp_dir = resume_helper.resume_gen.temp_dir
                json_path = os.path.join(temp_dir, "resume.json")
                
                profile = None
                if os.path.exists(json_path):
                    data, err = resume_helper.resume_gen.import_json(json_path)
                    if data is not None:
                        profile = data
                
                if profile is None:
                    return (
                        gr.update(value=None, interactive=False),
                        "❌ No resume data available. Please fill out your resume information first."
                    )
                
                profile.setdefault("full_name", profile.get("name", "Resume"))
                os.makedirs(temp_dir, exist_ok=True)
                
                full_name = profile.get("full_name", "Resume")
                safe_name = "".join(c for c in full_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                
                if format_choice == "pdf":
                    filename = f"{safe_name}_Resume.pdf" if safe_name else "Resume.pdf"
                    file_path = os.path.join(temp_dir, filename)
                    
                    ok = resume_helper.resume_gen.generate_pdf(profile, file_path)
                    if ok and os.path.exists(file_path):
                        embed_json_metadata_pdf(file_path, profile)
                        file_size = os.path.getsize(file_path)
                        abs_path = os.path.abspath(file_path)
                        return (
                            gr.update(value=abs_path, interactive=True),
                            f"✅ PDF resume generated successfully! ({file_size:,} bytes)"
                        )
                        
                elif format_choice == "json":
                    filename = f"{safe_name}_Resume.json" if safe_name else "Resume.json"
                    file_path = os.path.join(temp_dir, filename)
                    
                    ok, msg = resume_helper.resume_gen.export_json(profile, file_path)
                    if ok and os.path.exists(file_path):
                        file_size = os.path.getsize(file_path)
                        abs_path = os.path.abspath(file_path)
                        return (
                            gr.update(value=abs_path, interactive=True),
                            f"✅ JSON resume generated successfully! ({file_size:,} bytes)"
                        )
                        
                elif format_choice == "docx":
                    filename = f"{safe_name}_Resume.docx" if safe_name else "Resume.docx"
                    file_path = os.path.join(temp_dir, filename)
                    
                    ok = generate_docx_resume(profile, file_path)
                    if ok and os.path.exists(file_path):
                        file_size = os.path.getsize(file_path)
                        abs_path = os.path.abspath(file_path)
                        return (
                            gr.update(value=abs_path, interactive=True),
                            f"✅ Word document generated successfully! ({file_size:,} bytes)"
                        )
                    else:
                        return (
                            gr.update(value=None, interactive=False),
                            "❌ Word document generation failed. Please install python-docx: pip install python-docx"
                        )
                
                return (
                    gr.update(value=None, interactive=False),
                    f"❌ {format_choice.upper()} generation failed"
                )
                
            except Exception as e:
                return (
                    gr.update(value=None, interactive=False),
                    f"❌ Error generating resume: {str(e)}"
                )

        if all_tabs_components:
            personal = all_tabs_components.get("personal_info_tab", {})
            education = all_tabs_components.get("educations_tab", {})
            work = all_tabs_components.get("experiences_tab", {})
            skills = all_tabs_components.get("skills_tab", {})
            projects = all_tabs_components.get("projects_tab", {})
            certs = all_tabs_components.get("certifications_tab", {})
            others = all_tabs_components.get("others_tab", {})

            form_outputs = [
                personal.get(k) for k in [
                    "name_prefix_input","email_input","name_input","phone_input","current_address",
                    "location_input","citizenship","linkedin_input","github_input",
                    "portfolio_input","summary_input","info_output"
                ]
            ] + [
                education.get(k) for k in [
                    "institution_input","degree_input","field_input","edu_start_input",
                    "edu_end_input","gpa_input","edu_desc_input","edu_list"
                ]
            ] + [
                work.get(k) for k in [
                    "company_input","position_input","work_location_input","work_start_input",
                    "work_end_input","work_desc_input","achievements_input","work_list"
                ]
            ] + [
                skills.get(k) for k in [
                    "category_input","skill_input","proficiency_input","skill_list"
                ]
            ] + [
                projects.get(k) for k in [
                    "project_title_input","project_desc_input","project_tech_input",
                    "project_url_input","project_start_input","project_end_input",
                    "project_list"
                ]
            ] + [
                certs.get(k) for k in [
                    "cert_name_input","cert_issuer_input","cert_date_input",
                    "cert_id_input","cert_url_input","cert_list"
                ]
            ] + [
                others.get("sections_data") if others else gr.State({}),
                others.get("sections_display") if others else gr.Dataframe([]),
                others.get("section_selector") if others else gr.Dropdown([])
            ]

            def safe_file_handler(uploaded_file):
                """Basic file handler for Import & Export tab.
                
                Returns:
                    49 elements: status + 48 form outputs
                """
                try:
                    form_data, original_json = auto_load_file(uploaded_file)
                    return form_data
                except Exception:
                    empty_data = get_empty_form_data()
                    empty_data[0] = "❌ File processing error"
                    return empty_data

            def enhanced_file_handler(uploaded_file):
                """Enhanced file handler that also generates JSON display for AI Resume Helper tab.
                
                Returns:
                    50 elements: status + 48 form outputs + 1 JSON string
                """
                try:
                    form_data, original_json = auto_load_file(uploaded_file)
                    if form_data and len(form_data) > 0 and not form_data[0].startswith("❌"):
                        if original_json:
                            try:
                                temp_dir = resume_helper.resume_gen.temp_dir
                                import_cache_path = os.path.join(temp_dir, "last_import.json")
                                with open(import_cache_path, 'w', encoding='utf-8') as f:
                                    json.dump(original_json, f, indent=2)
                            except Exception:
                                pass
                            
                            json_data = json.dumps(original_json, indent=2)
                            return form_data + [json_data]
                        
                        try:
                            from utils.constants import UIConstants
                            from models.resume import ResumeSchema
                            
                            form_vals = form_data[1:]
                            
                            personal_end = UIConstants.PERSONAL_INFO_FIELDS
                            personal_fields = form_vals[0:personal_end]
                            
                            personal_field_order = ResumeSchema.get_field_order('personal_info')
                            personal_dict = {}
                            for i, field_name in enumerate(personal_field_order):
                                if i < len(personal_fields):
                                    personal_dict[field_name] = personal_fields[i]
                                else:
                                    personal_dict[field_name] = ""
                            
                            name_prefix = personal_dict.get('name_prefix', '')
                            email = personal_dict.get('email', '')
                            name = personal_dict.get('full_name', '')
                            phone = personal_dict.get('phone', '')
                            address = personal_dict.get('current_address', '')
                            location = personal_dict.get('location', '')
                            citizenship = personal_dict.get('citizenship', '')
                            linkedin = personal_dict.get('linkedin_url', '')
                            github = personal_dict.get('github_url', '')
                            portfolio = personal_dict.get('portfolio_url', '')
                            summary = personal_dict.get('summary', '')
                            
                            edu_table = form_vals[UIConstants.EDUCATION_TABLE_INDEX - 1]
                            work_table = form_vals[UIConstants.WORK_TABLE_INDEX - 1]
                            skill_table = form_vals[UIConstants.SKILLS_TABLE_INDEX - 1]
                            project_table = form_vals[UIConstants.PROJECTS_TABLE_INDEX - 1]
                            cert_table = form_vals[UIConstants.CERTIFICATIONS_TABLE_INDEX - 1]
                            
                            others_data = form_vals[UIConstants.OTHERS_SECTIONS_DATA_INDEX - 1]
                            
                            profile = resume_helper.build_profile_dict(
                                name_prefix, email, name, phone, address, location, citizenship,
                                linkedin, github, portfolio, summary,
                                edu_table, work_table, skill_table, project_table, cert_table,
                                others_data
                            )
                            
                            json_data = json.dumps(profile, indent=2)
                            return form_data + [json_data]
                        except Exception as e:
                            from utils.logging_helpers import log_error
                            log_error(f"Error reconstructing JSON from form data: {e}")
                            return form_data + [f"Error generating JSON: {str(e)}"]
                    else:
                        return form_data + [""]
                        
                except Exception as e:
                    from utils.logging_helpers import log_error
                    log_error(f"Error in enhanced_file_handler: {e}")
                    empty_data = get_empty_form_data()
                    empty_data[0] = "❌ File processing error"
                    return empty_data + [""]
            file_import.upload(
                safe_file_handler, 
                inputs=[file_import], 
                outputs=[import_status] + form_outputs
            )
            
            def clear_all_fields():
                return get_empty_form_data()
                
            file_import.clear(
                clear_all_fields,
                inputs=[],
                outputs=[import_status] + form_outputs
            )

            def generate_with_update(format_choice, name_prefix, email, name, phone, current_address, location, citizenship,
                                   linkedin_url, github_url, portfolio_url, summary, edu_table, experience_table, 
                                   skill_table, project_table, certifications_table, others_data=None):
                try:
                    profile = resume_helper.build_profile_dict(
                        name_prefix, email, name, phone, current_address, location, citizenship,
                        linkedin_url, github_url, portfolio_url, summary,
                        edu_table, experience_table, skill_table, project_table, certifications_table,
                        others_data
                    )
                    
                    temp_dir = resume_helper.resume_gen.temp_dir
                    os.makedirs(temp_dir, exist_ok=True)
                    json_path = os.path.join(temp_dir, "resume.json")
                    
                    with open(json_path, 'w', encoding='utf-8') as f:
                        json.dump(profile, f, indent=2)
                    
                    return generate_resume(format_choice)
                    
                except Exception as e:
                    return (
                        gr.update(value=None, interactive=False),
                        f"❌ Error updating data: {str(e)}"
                    )
            others_tab = all_tabs_components.get("others_tab", {})
            generation_inputs = [
                export_format,
                personal.get("name_prefix_input"),
                personal.get("email_input"),   personal.get("name_input"),
                personal.get("phone_input"),   personal.get("current_address"),
                personal.get("location_input"),personal.get("citizenship"),
                personal.get("linkedin_input"),personal.get("github_input"),
                personal.get("portfolio_input"),personal.get("summary_input"),
                education.get("edu_list"), work.get("work_list"),
                skills.get("skill_list"),  projects.get("project_list"),
                certs.get("cert_list"),
                others_tab.get("sections_data") if others_tab else gr.State({}),
            ]

            generate_btn.click(
                generate_with_update,
                inputs=generation_inputs,
                outputs=[download_btn, export_status]
            )

    def connect_to_ai_resume_helper(ai_helper_tab):
        if ai_helper_tab and hasattr(ai_helper_tab, 'components') and 'resume_json' in ai_helper_tab.components:
            ai_resume_json = ai_helper_tab.components['resume_json']
            
            file_import.upload(
                enhanced_file_handler, 
                inputs=[file_import], 
                outputs=[import_status] + form_outputs + [ai_resume_json]
            )
            
            def clear_all_fields_with_json():
                """Clear all fields and reset AI resume JSON.
                
                Returns: status + form_outputs (48) + ai_resume_json (1) = 50 total
                Note: get_empty_form_data() already includes status, so we use it for form_outputs only
                """
                empty_data = get_empty_form_data()
                return empty_data + [""]
            
            file_import.clear(
                clear_all_fields_with_json,
                inputs=[],
                outputs=[import_status] + form_outputs + [ai_resume_json]
            )
            
            return True
        else:
            return False
    return SimpleNamespace(
        tab=tab,
        components={
            "file_import": file_import,
            "export_format": export_format,
            "generate_btn": generate_btn,
            "download_btn": download_btn,
            "import_status": import_status,
            "export_status": export_status,
        },
        connect_to_ai_resume_helper=connect_to_ai_resume_helper,
    )
